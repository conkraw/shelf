import streamlit as st
import pandas as pd
import os
import glob
import random
import datetime
import re
from dateutil import tz

from docx import Document
from docx.shared import Inches

import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

import firebase_admin
from firebase_admin import credentials, firestore

# Set wide layout
st.set_page_config(layout="wide")

# Initialize Firebase if not already done.
firebase_creds = st.secrets["firebase_service_account"].to_dict()
if not firebase_admin._apps:
    cred = credentials.Certificate(firebase_creds)
    firebase_admin.initialize_app(cred)
db = firestore.client()

### Helper functions to manage exam state in Firestore

def initialize_state():
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False
    if "score" not in st.session_state:
        st.session_state.score = 0
    if "question_index" not in st.session_state:
        st.session_state.question_index = 0
    if "results" not in st.session_state:
        st.session_state.results = []
    if "selected_answers" not in st.session_state:
        st.session_state.selected_answers = []
    if "user_name" not in st.session_state:
        st.session_state.user_name = ""
    if "assigned_passcode" not in st.session_state:
        st.session_state.assigned_passcode = ""
    if "recipient_email" not in st.session_state:
        st.session_state.recipient_email = ""
    if "df" not in st.session_state:
        st.session_state.df = None
    if "result_message" not in st.session_state:
        st.session_state.result_message = ""
    if "result_color" not in st.session_state:
        st.session_state.result_color = ""
    if "result_messages" not in st.session_state:
        st.session_state.result_messages = []
    if "question_ids" not in st.session_state:
        st.session_state.question_ids = []

def get_user_key():
    # Use the entire assigned passcode as the key.
    return str(st.session_state.assigned_passcode)

def save_exam_state():
    user_key = get_user_key()
    data = {
        "question_index": st.session_state.question_index,
        "score": st.session_state.score,
        "results": st.session_state.results,
        "selected_answers": st.session_state.selected_answers,
        "result_messages": st.session_state.result_messages,
        "question_ids": st.session_state.question_ids,
        "email_sent": st.session_state.get("email_sent", False),
        "exam_complete": st.session_state.get("exam_complete", False), 
        "timestamp": firestore.SERVER_TIMESTAMP,
    }
    db.collection("exam_sessions").document(user_key).set(data)

def load_exam_state():
    user_key = get_user_key()
    doc_ref = db.collection("exam_sessions").document(user_key)
    doc = doc_ref.get()
    if doc.exists:
        data = doc.to_dict()
        st.session_state.question_index = data.get("question_index", 0)
        st.session_state.score = data.get("score", 0)
        st.session_state.results = data.get("results", st.session_state.results)
        st.session_state.selected_answers = data.get("selected_answers", st.session_state.selected_answers)
        st.session_state.result_messages = data.get("result_messages", st.session_state.result_messages)
        st.session_state.question_ids = data.get("question_ids", st.session_state.question_ids)
        st.session_state.email_sent = data.get("email_sent", False)

def create_new_exam(full_df):
    """
    Samples 5 questions from full_df and initializes the exam state.
    """
    sample_df = sample_new_exam(full_df, n=5)
    
    st.session_state.question_ids = list(sample_df["record_id"])
    st.session_state.df = sample_df.reset_index(drop=True)
    total_questions = len(st.session_state.df)
    st.session_state.results = [None] * total_questions
    st.session_state.selected_answers = [None] * total_questions
    st.session_state.result_messages = ["" for _ in range(total_questions)]
    
def check_and_add_passcode(passcode):
    passcode_str = str(passcode)
    if passcode_str.lower() == "password":
        return False
    doc_ref = db.collection("shelf_records").document(passcode_str)
    if not doc_ref.get().exists:
        doc_ref.set({"processed": True})
        return False
    else:
        return True



def is_passcode_locked(passcode, lock_hours=6):
    """
    Checks if the passcode is locked.
    Returns True if locked (i.e. the passcode was locked within the last lock_hours),
    otherwise returns False.
    """
    doc_ref = db.collection("locked_passcodes").document(str(passcode))
    doc = doc_ref.get()
    if doc.exists:
        data = doc.to_dict()
        lock_time = data.get("lock_time")
        if lock_time is not None:
            # lock_time is already a timezone-aware datetime.
            now = datetime.datetime.now(datetime.timezone.utc)
            delta = now - lock_time
            if delta.total_seconds() < lock_hours * 3600:
                return True
    return False


def lock_passcode(passcode):
    """
    Locks the passcode by writing the current server timestamp to Firestore.
    This marks the passcode as used and locked for 6 hours.
    """
    doc_ref = db.collection("locked_passcodes").document(str(passcode))
    # Set the lock time to the server timestamp.
    doc_ref.set({"lock_time": firestore.SERVER_TIMESTAMP})

def get_or_set_passcode_start(passcode):
    ref = db.collection("passcode_starts").document(passcode)
    doc = ref.get()
    if doc.exists:
        # Firestore returns a timezone-aware UTC datetime
        return doc.to_dict()["start_time"]
    else:
        ref.set({"start_time": firestore.SERVER_TIMESTAMP})
        # until Firestore syncs back, approximate with local now in UTC
        return datetime.datetime.now(datetime.timezone.utc)

def passcode_expires_at(start_utc: datetime.datetime) -> datetime.datetime:
    """
    Expires on the same week's Friday at 23:59:59 local time (America/New_York).
    Converts the UTC timestamp into local, finds that Friday date,
    sets 23:59:59 local, then returns an equivalent UTC datetime.
    """
    # 1) Define & convert into local Eastern Time
    LOCAL_TZ    = tz.gettz("America/New_York")
    start_local = start_utc.astimezone(LOCAL_TZ)

    # 2) Compute that week's Friday (weekday 4)
    base    = start_local.date()
    wd      = base.weekday()  # Mon=0 … Sun=6
    if wd <= 4:
        days_to_fri = 4 - wd
    else:
        days_to_fri = 4 + 7 - wd
    fri_date = base + datetime.timedelta(days=days_to_fri)

    # 3) Build a local datetime at 23:59:59
    expiry_local = datetime.datetime.combine(
        fri_date,
        datetime.time(hour=23, minute=59, second=59),
        tzinfo=LOCAL_TZ
    )

    # 4) Convert back to UTC for your is_passcode_expired check
    return expiry_local.astimezone(datetime.timezone.utc)

def is_passcode_expired(passcode: str) -> bool:
    start_utc = get_or_set_passcode_start(passcode)
    expiry_utc = passcode_expires_at(start_utc)
    now_utc    = datetime.datetime.now(datetime.timezone.utc)
    return now_utc > expiry_utc
    
def get_image_path(record_id, folder="images"):
    extensions = ["jpg", "jpeg", "png", "gif"]
    for ext in extensions:
        pattern = os.path.join(folder, f"{record_id}.{ext}")
        matches = glob.glob(pattern)
        if matches:
            return matches[0]
    return None

def get_global_used_questions():
    """
    Retrieves a list of question record_ids that have been used in the last 7 days.
    Documents older than 7 days are deleted so questions can be reused.
    """
    used_questions_ref = db.collection("global_used_questions")
    docs = used_questions_ref.stream()
    used_ids = []
    now = datetime.datetime.utcnow()
    for doc in docs:
        data = doc.to_dict()
        ts = data.get("timestamp")
        if ts is not None:
            ts_naive = ts.replace(tzinfo=None)
            if (now - ts_naive).days < 7:
                used_ids.append(doc.id)
            else:
                # Delete outdated documents so questions become available.
                doc.reference.delete()
    return used_ids

def mark_questions_as_used(question_ids):
    """
    Marks the given list of question_ids as used globally,
    storing a timestamp so that they can be reset after 7 days.
    """
    used_questions_ref = db.collection("global_used_questions")
    for qid in question_ids:
        used_questions_ref.document(str(qid)).set({
            "used": True,
            "timestamp": firestore.SERVER_TIMESTAMP
        })

def sample_new_exam(full_df, n=5):
    """
    Samples n questions from full_df that have not yet been used in the last 7 days.
    If no questions are available, displays an error message and stops.
    If fewer than n questions are available, uses all remaining questions.
    """
    used_ids = get_global_used_questions()
    available_df = full_df[~full_df["record_id"].isin(used_ids)]
    if available_df.empty:
        st.error("No further cases available for your passcode. Please try again later.")
        st.stop()
    if len(available_df) < n:
        st.warning("Fewer than the expected number of questions are available. Using all remaining questions.")
        sample_df = available_df
    else:
        sample_df = available_df.sample(n=n, replace=False)
    mark_questions_as_used(sample_df["record_id"].tolist())
    return sample_df


def load_data(pattern="*.csv"):
    csv_files = glob.glob(pattern)
    dfs = [pd.read_csv(file) for file in csv_files]
    combined_df = pd.concat(dfs, ignore_index=True)
    if "record_id" not in combined_df.columns:
        combined_df["record_id"] = combined_df.index + 1
    return combined_df
    
def generate_review_doc(row, user_selected_letter, output_filename="review.docx"):
    doc = Document()
    doc.add_heading("Review of Incorrect Question", level=1)
    doc.add_heading(f"Student: {st.session_state.user_name}", level=2)
    doc.add_heading(f"Question {row['record_id']}:", level=2)
    doc.add_paragraph(row["question"])
    
    image_path = get_image_path(row["record_id"])
    if image_path:
        try:
            doc.add_picture(image_path, width=Inches(4))
        except Exception as e:
            doc.add_paragraph(f"(Image could not be added: {e})")
    
    if "anchor" in row:
        doc.add_paragraph(row["anchor"])
    
    doc.add_heading("Answer Choices:", level=2)
    for letter in ["a", "b", "c", "d", "e"]:
        col_name = "answerchoice_" + letter
        if pd.notna(row[col_name]) and str(row[col_name]).strip():
            doc.add_paragraph(f"{letter.upper()}: {row[col_name]}")
    
    doc.add_heading("Student Answer:", level=2)
    if user_selected_letter:
        user_answer_text = row.get("answerchoice_" + user_selected_letter, "N/A")
        doc.add_paragraph(user_answer_text)
    else:
        doc.add_paragraph("No answer selected.")
    
    correct_letter = str(row["correct_answer"]).strip().lower()
    correct_answer_text = row.get("answerchoice_" + correct_letter, "N/A")
    doc.add_heading("Correct Answer:", level=2)
    doc.add_paragraph(correct_answer_text)
    
    doc.add_heading("Explanation:", level=2)
    doc.add_paragraph(row["answer_explanation"])
    
    doc.save(output_filename)
    return output_filename

def send_email_with_attachment(to_emails, subject, body, attachment_path):
    from_email = st.secrets["general"]["email"]
    password = st.secrets["general"]["email_password"]
    
    msg = MIMEMultipart()
    msg['From'] = from_email
    msg['To'] = ', '.join(to_emails)
    msg['Subject'] = subject
    msg.attach(MIMEText(body, 'html'))
    
    with open(attachment_path, 'rb') as attachment:
        part = MIMEBase('application', 'octet-stream')
        part.set_payload(attachment.read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', 'attachment', filename=os.path.basename(attachment_path))
        msg.attach(part)
    
    try:
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
            server.login(from_email, password)
            server.send_message(msg, from_addr=from_email, to_addrs=to_emails)
        st.success("Email sent successfully!")
    except Exception as e:
        st.error(f"Error sending email: {e}")

def save_exam_results():
    """
    Collects exam results details and saves them to the 'exam_results' collection in Firestore.
    The details include for each question:
      - record_id
      - student's answer (the letter)
      - correct answer text
      - a result flag ("Correct" or "Incorrect")
    It also saves the student's name, the passcode used, and the overall score.
    """
    exam_data = []
    df = st.session_state.df  # This is the exam DataFrame for this session.
    # Iterate over each question in the exam.
    for idx, row in df.iterrows():
        record = {}
        record["record_id"] = row["record_id"]
        
        # Get the student's answer for this question.
        student_ans = st.session_state.selected_answers[idx]
        record["student_answer"] = student_ans if student_ans is not None else ""
        
        # Determine correct answer: we assume your DataFrame has a "correct_answer" field,
        # and answer choices are stored in columns like "answerchoice_a", "answerchoice_b", etc.
        correct_letter = str(row["correct_answer"]).strip().lower()
        correct_answer_text = row.get("answerchoice_" + correct_letter, "")
        record["correct_answer"] = correct_answer_text
        
        # Set result.
        if student_ans and student_ans == correct_letter:
            record["result"] = "Correct"
        else:
            record["result"] = "Incorrect"
        
        exam_data.append(record)
    
    # Prepare a summary dictionary.
    exam_summary = {
        "student_name": st.session_state.user_name,
        "passcode": st.session_state.assigned_passcode,
        "score": st.session_state.score,
        "total_questions": len(df),
        "exam_data": exam_data,
        "timestamp": firestore.SERVER_TIMESTAMP,
    }
    
    # Save to the "exam_results" collection.
    db.collection("exam_results").add(exam_summary)
    st.success("Thank you for your participation!")
    
### Login Screen

def login_screen():
    st.title("Pediatric Clerkship NBME-Style Assessment Portal")
    
    with st.expander("📖Instructions", expanded=True):
        col1, col2 = st.columns(2)
    
        with col1:
            st.markdown("""
            1. Each student will be provided a **password** by their preceptor.  
               This grants you access to **five** NBME‑style questions in various pediatric subjects, each including an answer and explanation.
    
            2. If you answer one or more questions **incorrectly**, a random incorrect question  
               (with your answer and the explanation) will be emailed to the preceptor who gave you the password,  
               so they can review it with you.
    
            3. While every effort has been made for accuracy, **errors can occur**.  
               If you believe a question or explanation is in error, please email the **course director** for review.
            """)
    
        with col2:
            st.markdown("""
            4. After you **complete** your set of 5 questions, your passcode will be **locked for 6 hours**  
               to prevent multiple emails.
    
            5. **No question will be repeated** during the week your passcode is in use.
    
            6. **All passcodes expire by Friday at 23:59** of the week in which they are first used.  
               - *For example:* if you first use your password on **Sunday, April 6**, it will expire at **Friday, April 11 23:59**.  
               - If you first use it on **Friday morning**, it still expires that same **Friday at 23:59**.
            """)

    passcode_input = st.text_input("Enter your assigned passcode", type="password")
    user_name = st.text_input("Enter your name")
    
    if st.button("Login"):
        if "recipients" not in st.secrets:
            st.error("Recipient emails not configured. Please set them in your secrets file under [recipients].")
            return
        if passcode_input not in st.secrets["recipients"]:
            st.error("Invalid passcode. Please try again.")
            return
        if not user_name:
            st.error("Please enter your name to proceed.")
            return

        if is_passcode_expired(passcode_input):
            st.error("This passcode has expired for the week. Contact your instructor.")
            return
    
        # Load the full dataset from CSVs.
        full_df = load_data()  # Loads all CSV files.
        
        # Optionally filter by subject based on designation in the passcode.
        subject_mapping = {
            "aaa": "Respiratory",
            "aab": "School-Based",
            # add more mappings as needed...
        }
        if "_" in passcode_input:
            designation = passcode_input.split("_")[-1]  # get part after underscore
            if designation in subject_mapping:
                subject_filter = subject_mapping[designation]
                filtered_df = full_df[full_df["subject"] == subject_filter]
                if not filtered_df.empty:
                    full_df = filtered_df
                else:
                    st.warning(f"No questions found for subject {subject_filter}. Using full dataset instead.")
        
        # Check for a saved exam session.
        #user_key = str(st.session_state.assigned_passcode)
        user_key = passcode_input
        doc_ref = db.collection("exam_sessions").document(user_key)
        doc = doc_ref.get()

        if doc.exists:
            data = doc.to_dict()
            # Check if the saved session is complete.
            if data.get("exam_complete", False):
                # Exam was complete; now check if the lock period is still active.
                if is_passcode_locked(passcode_input, lock_hours=6):
                    st.error("This passcode is locked for 6 hours. Please try again later.")
                    return
                else:
                    # Lock period has expired—delete the old session and create a new exam.
                    doc_ref.delete()
                    create_new_exam(full_df)
            else:
                # Resume the incomplete exam session.
                st.session_state.question_index = data.get("question_index", 0)
                st.session_state.score = data.get("score", 0)
                st.session_state.results = data.get("results", [])
                st.session_state.selected_answers = data.get("selected_answers", [])
                st.session_state.result_messages = data.get("result_messages", [])
                st.session_state.question_ids = data.get("question_ids", [])
                if st.session_state.question_ids:
                    qids = st.session_state.question_ids
                    sample_df = full_df[full_df["record_id"].isin(qids)]
                    sample_df = sample_df.set_index("record_id").loc[qids].reset_index()
                    st.session_state.df = sample_df
                else:
                    st.session_state.df = full_df
        else:
            # No saved session exists: create a new exam.
            create_new_exam(full_df)


        # Save the login details in session state.
        st.session_state.assigned_passcode = passcode_input
        st.session_state.recipient_email = st.secrets["recipients"][passcode_input]
        st.session_state.user_name = user_name
        st.session_state.authenticated = True
        
        st.rerun()

### Exam Screen

def exam_screen():
    st.title("Pediatric Clerkship NBME-Style Assessment Portal")
    st.write(f"Welcome, **{st.session_state.user_name}**!")
    
    df = st.session_state.df
    total_questions = len(df)
    
    with st.sidebar:
        st.header("Navigation")
        for i in range(total_questions):
            marker = ""
            if st.session_state.results[i] == "correct":
                marker = "✅"
            elif st.session_state.results[i] == "incorrect":
                marker = "❌"
            current_marker = " (Current)" if i == st.session_state.question_index else ""
            label = f"Question {i+1}: {marker}{current_marker}"
            if st.button(label, key=f"nav_{i}"):
                st.session_state.question_index = i
                st.rerun()
    
    if st.session_state.question_index >= total_questions:
        percentage = (st.session_state.score / total_questions) * 100
        st.header("Exam Completed")
        st.write(f"Your final score is **{st.session_state.score}** out of **{total_questions}** ({percentage:.1f}%).")
        
        # Mark the exam as complete.
        st.session_state.exam_complete = True
        save_exam_state()  # Save the complete state.
        
        # Lock the passcode if not already locked.
        if not is_passcode_locked(st.session_state.assigned_passcode, lock_hours=6):
            lock_passcode(st.session_state.assigned_passcode)
            st.success("Your passcode has now been locked for 6 hours and cannot be used again.")
        
        # Send review email only once.
        if not st.session_state.get("email_sent", False):
            wrong_indices = [i for i, result in enumerate(st.session_state.results) if result == "incorrect"]
            if wrong_indices:
                selected_index = random.choice(wrong_indices)
                selected_row = st.session_state.df.iloc[selected_index]
                doc_filename = f"review_{st.session_state.user_name}_q{selected_index+1}.docx"
                generate_review_doc(selected_row, st.session_state.selected_answers[selected_index], output_filename=doc_filename)
                try:
                    send_email_with_attachment(
                        to_emails=[st.session_state.recipient_email],
                        subject="Review of an Incorrect Question",
                        body="Please find attached a review document for a question answered incorrectly.",
                        attachment_path=doc_filename
                    )
                    #st.success("Review email sent successfully!")
                    st.session_state.email_sent = True
                    save_exam_state()
                except Exception as e:
                    st.error(f"Error sending email: {e}")
            else:
                st.info("No incorrect answers to review!")
        else:
            st.info("Review email has already been sent for this exam.")
        
        save_exam_results()
        return

    # Get the current row
    current_row = df.iloc[st.session_state.question_index]
    
    # Check if the current row has the expected keys. For example, verify "answerchoice_a" exists.
    if "answerchoice_a" not in current_row:
        st.error("No further questions available for your exam. Please try again later.")
        st.stop()
    
    # Proceed as usual:
    option_cols = [
        ("a", current_row["answerchoice_a"]),
        ("b", current_row["answerchoice_b"]),
        ("c", current_row["answerchoice_c"]),
        ("d", current_row["answerchoice_d"]),
        ("e", current_row["answerchoice_e"]),
    ]
    options = []
    option_mapping = {}
    for letter, text in option_cols:
        if pd.notna(text) and str(text).strip():
            option_text = f"{letter.upper()}. {text.strip()}"
            options.append(option_text)
            option_mapping[option_text] = letter

    answered = st.session_state.selected_answers[st.session_state.question_index] is not None
    default_index = 0
    if answered:
        selected_letter = st.session_state.selected_answers[st.session_state.question_index]
        default_option = None
        for opt, letter in option_mapping.items():
            if letter == selected_letter:
                default_option = opt
                break
        if default_option in options:
            default_index = options.index(default_option)
    
    col1, col2 = st.columns(2)
    with col1:
        st.write(f"**Question ({current_row['record_id']}):**")
        st.write(current_row["question"])
        record_id = current_row["record_id"]
        image_path = get_image_path(record_id)
        if image_path:
            st.image(image_path, use_container_width=True)
        st.write(current_row["anchor"])
        st.write("**Select your answer:**")
        answer_text_mapping = {}
        letter_to_answer = {}
        options = []
        for letter in ["a", "b", "c", "d", "e"]:
            col_name = "answerchoice_" + letter
            if pd.notna(current_row[col_name]) and str(current_row[col_name]).strip():
                text = str(current_row[col_name]).strip()
                options.append(text)
                answer_text_mapping[text] = letter
                letter_to_answer[letter] = text
        for i, option in enumerate(options):
            if not answered:
                if st.button(option, key=f"option_{st.session_state.question_index}_{i}"):
                    selected_letter = answer_text_mapping[option]
                    st.session_state.selected_answers[st.session_state.question_index] = selected_letter
                    correct_answer_letter = str(current_row["correct_answer"]).strip().lower()
                    if selected_letter == correct_answer_letter:
                        st.session_state.results[st.session_state.question_index] = "correct"
                        message = "Correct!"
                        st.session_state.score += 1
                    else:
                        st.session_state.results[st.session_state.question_index] = "incorrect"
                        correct_answer_text = letter_to_answer.get(correct_answer_letter, "")
                        message = f"Incorrect. The correct answer was: {correct_answer_text}"
                    
                    st.session_state.result_messages[st.session_state.question_index] = message
                    
                    save_exam_state()
                    st.rerun()
            else:
                st.button(option, key=f"option_{st.session_state.question_index}_{i}", disabled=True)
    
    with col2:
        if answered:
            result_msg = st.session_state.result_messages[st.session_state.question_index]
            if st.session_state.results[st.session_state.question_index] == "correct":
                st.success(result_msg)
            elif st.session_state.results[st.session_state.question_index] == "incorrect":
                st.error(result_msg)
            
            st.write("**Explanation:**")
            st.write(current_row["answer_explanation"])
            
            if st.button("Next Question"):
                st.session_state.question_index += 1
                st.session_state.result_message = ""
                st.session_state.result_color = ""
                save_exam_state()
                st.rerun()


def main():
    initialize_state()
    if not st.session_state.authenticated:
        login_screen()
    else:
        exam_screen()
        
if __name__ == "__main__":
    main()

